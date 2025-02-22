import argparse
import re
import os
from docx import Document

def replace_text(text, old_pattern, new_text, ignore_case=False):
    flags = re.IGNORECASE if ignore_case else 0
    return re.sub(old_pattern, new_text, text, flags=flags)

def process_paragraphs(doc, replacements):
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        for old, new, ignore_case in replacements:
            original_text = replace_text(original_text, re.escape(old), new, ignore_case)
        paragraph.text = original_text

def process_tables(doc, replacements):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                for old, new, ignore_case in replacements:
                    original_text = replace_text(original_text, re.escape(old), new, ignore_case)
                cell.text = original_text

def replace_in_docx(input_path, output_path):
    # Define the replacements: (old_text, new_text, ignore_case flag)
    replacements = [
        ("Ashish Mishra", "Abdurrahman Sarguroh", True),
        ("122AX016", "122AX002", False)
    ]

    doc = Document(input_path)
    process_paragraphs(doc, replacements)
    process_tables(doc, replacements)
    doc.save(output_path)
    print(f"Replaced text in '{input_path}' and saved the updated document as '{output_path}'.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Replace names and roll numbers in a Word document.")
    parser.add_argument("input_file", help="Path to the input .docx file")
    args = parser.parse_args()
    
    # Generate output file name by replacing "122AX016" with "122AX002" in the input file name.
    base_name = os.path.basename(args.input_file)
    new_base_name = base_name.replace("122AX016", "122AX002")
    dir_name = os.path.dirname(args.input_file)
    output_path = os.path.join(dir_name, new_base_name)
    
    replace_in_docx(args.input_file, output_path)
